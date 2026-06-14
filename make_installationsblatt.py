from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "installationsblatt"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "KLARA_Installations-und-Datensicherungsblatt.docx"

BLUE = "0879BD"
BLUE_DARK = "075E91"
BLUE_SOFT = "EAF8FD"
CYAN_SOFT = "EDF9FA"
GREEN = "177A55"
GREEN_SOFT = "EAF7F1"
AMBER = "9A6200"
AMBER_SOFT = "FFF5DA"
RED = "9A304B"
RED_SOFT = "FFF0F3"
INK = "19313F"
MUTED = "587080"
LINE = "D7E5ED"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text="", size=10.5, bold=False, color=INK, after=5, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    set_run(run, size=15 if level == 1 else 12, bold=True, color=BLUE_DARK)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def add_step(doc, number, title, detail, accent=BLUE, fill=BLUE_SOFT):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(1.15)
    table.columns[1].width = Cm(14.8)
    left, right = table.rows[0].cells
    left.width = Cm(1.15)
    right.width = Cm(14.8)
    for cell in (left, right):
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=fill, size="2")
        set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_p.paragraph_format.space_after = Pt(0)
    set_run(left_p.add_run(str(number)), size=16, bold=True, color=accent)
    right_p = right.paragraphs[0]
    right_p.paragraph_format.space_after = Pt(1)
    set_run(right_p.add_run(title), size=10.5, bold=True, color=INK)
    detail_p = right.add_paragraph()
    detail_p.paragraph_format.space_after = Pt(0)
    detail_p.paragraph_format.line_spacing = 1.1
    set_run(detail_p.add_run(detail), size=9.5, color=MUTED)
    add_text(doc, "", after=2)


def add_callout(doc, title, body, fill, accent):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="10")
    set_cell_margins(cell, top=160, start=200, bottom=150, end=200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run(p2.add_run(body), size=10, color=INK)
    add_text(doc, "", after=2)


def add_page_header(doc, kicker, title, subtitle):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(11.8)
    logo_cell, title_cell = table.rows[0].cells
    for cell in (logo_cell, title_cell):
        set_cell_border(cell, color=WHITE, size="0")
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.add_run().add_picture(str(ROOT / "logo.png"), width=Cm(3.7))
    p = title_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(kicker.upper()), size=8, bold=True, color=BLUE)
    p2 = title_cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    set_run(p2.add_run(title), size=21, bold=True, color=BLUE_DARK)
    p3 = title_cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    set_run(p3.add_run(subtitle), size=9.5, color=MUTED)
    add_text(doc, "", after=4)


def add_device_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.4), Cm(5.8), Cm(6.8)]
    headers = ["Gerät", "Empfohlener Browser", "Installation"]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        set_cell_shading(cell, BLUE)
        set_cell_border(cell, color=BLUE)
        set_cell_margins(cell)
        set_run(cell.paragraphs[0].add_run(headers[i]), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("iPhone / iPad", "Safari", "Teilen-Symbol antippen, „Zum Home-Bildschirm“ wählen, dann „Hinzufügen“."),
        ("Android", "Google Chrome", "Menü ⋮ öffnen, „App installieren“ oder „Zum Startbildschirm hinzufügen“ wählen."),
        ("Windows / macOS", "Chrome oder Edge", "Installationssymbol in der Adresszeile oder Browsermenü öffnen und „App installieren“ wählen."),
    ]
    for row_i, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].width = widths[i]
            set_cell_shading(cells[i], WHITE if row_i % 2 == 0 else "F7FBFD")
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            set_run(p.add_run(value), size=9.2, bold=(i == 0), color=INK)
    add_text(doc, "", after=2)


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("KLARA · Installations- und Datensicherungsblatt · Stand 14.06.2026"), size=8, color=MUTED)


def set_picture_alt_text(document, description):
    for shape in document.inline_shapes:
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", description)
        doc_pr.set("title", description)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.25)
section.bottom_margin = Cm(1.25)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)
section.header_distance = Cm(0.6)
section.footer_distance = Cm(0.6)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15

for level, size in ((1, 15), (2, 12)):
    style = styles[f"Heading {level}"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE_DARK)

add_footer(section)

add_page_header(
    doc,
    "KLARA Begleitblatt",
    "KLARA als App installieren",
    "Anleitung für Smartphone, Tablet und Computer · mit Hinweisen zur Datensicherung",
)
add_callout(
    doc,
    "Wichtig vor dem ersten Eintrag",
    "Installieren Sie KLARA möglichst zuerst und öffnen Sie die Anwendung danach immer über dasselbe App-Symbol. "
    "Je nach Gerät und Browser können die Website und die installierte App getrennte lokale Datenbestände verwenden.",
    AMBER_SOFT,
    AMBER,
)

add_heading(doc, "1. App-Adresse öffnen", 1)
add_text(doc, "Öffnen Sie den offiziellen Link, den Sie von Ihrer Fachperson oder Einrichtung erhalten haben.")
add_callout(
    doc,
    "Offizielle App-Adresse",
    "____________________________________________________________\n"
    "Verwenden Sie nur den von der Einrichtung bereitgestellten Link.",
    CYAN_SOFT,
    BLUE_DARK,
)

add_heading(doc, "2. Auf dem Smartphone oder Tablet installieren", 1)
add_device_table(doc)
add_text(
    doc,
    "Hinweis: Die Bezeichnungen können je nach Browser- und Betriebssystemversion leicht abweichen.",
    size=8.5,
    color=MUTED,
    italic=True,
    after=6,
)

add_heading(doc, "3. Nach der Installation", 1)
add_bullet(doc, "Öffnen Sie KLARA künftig über das App-Symbol auf dem Start- oder Home-Bildschirm.")
add_bullet(doc, "Prüfen Sie nach dem ersten Öffnen, ob KLARA auch ohne Internetverbindung startet.")
add_bullet(doc, "Löschen Sie die App, Browserdaten oder Website-Daten nicht, solange keine aktuelle Sicherung vorhanden ist.")

doc.add_page_break()
add_page_header(
    doc,
    "KLARA Begleitblatt",
    "Auf mehreren Geräten nutzen",
    "Keine automatische Synchronisierung · sichere Übertragung per Export und Import",
)
add_callout(
    doc,
    "Keine automatische Synchronisierung",
    "KLARA speichert Einträge ausschließlich lokal auf dem jeweils verwendeten Gerät und im verwendeten Browser. "
    "Einträge auf Smartphone, Tablet und Computer werden aus Datenschutzgründen nicht automatisch zusammengeführt.",
    BLUE_SOFT,
    BLUE_DARK,
)

add_heading(doc, "Daten auf ein anderes Gerät übertragen", 1)
add_step(doc, 1, "Auf dem bisherigen Gerät exportieren", "In KLARA „Mehr“ öffnen und unter „Daten sichern und übertragen“ auf „Sicherung exportieren“ tippen.")
add_step(doc, 2, "Sicherungsdatei geschützt übertragen", "Die Datei enthält sensible Gesundheitsdaten im Klartext. Nutzen Sie einen geschützten Übertragungsweg und senden Sie sie nicht unverschlüsselt an beliebige Empfänger.")
add_step(doc, 3, "KLARA auf dem neuen Gerät installieren", "Installieren Sie KLARA mit der Anleitung auf Seite 1 und öffnen Sie sie über das neue App-Symbol.")
add_step(doc, 4, "Sicherung importieren", "Unter „Mehr“ auf „Sicherung importieren“ tippen und die exportierte Datei auswählen. Ein Import ersetzt nach Bestätigung die aktuell auf diesem Gerät gespeicherten Daten.")
add_step(doc, 5, "Übertragung kontrollieren", "Prüfen Sie einige Einträge und erstellen Sie anschließend auf dem neuen Gerät eine aktuelle Sicherung.")

add_heading(doc, "Empfohlene Nutzung", 1)
add_bullet(doc, "Führen Sie das Tagebuch möglichst auf einem Hauptgerät.")
add_bullet(doc, "Wenn Sie mehrere Geräte nutzen, exportieren Sie vor jedem Gerätewechsel die neueste Sicherung.")
add_bullet(doc, "Vermeiden Sie parallele Einträge auf mehreren Geräten: Beim Import werden Daten ersetzt, nicht automatisch zusammengeführt.")

doc.add_page_break()
add_page_header(
    doc,
    "KLARA Begleitblatt",
    "Schützt die Installation vor Datenverlust?",
    "Kurze Antwort: nur teilweise · regelmäßige Sicherungen bleiben notwendig",
)
add_callout(
    doc,
    "Die Installation allein verhindert keinen Datenverlust.",
    "Die App-Installation erleichtert den Zugriff, ermöglicht die Offline-Nutzung und kann versehentliches Arbeiten im falschen Browserfenster reduzieren. "
    "Die Einträge bleiben jedoch lokale Browser- beziehungsweise App-Daten und sind keine gesicherte Cloud-Kopie.",
    RED_SOFT,
    RED,
)

add_heading(doc, "Daten können weiterhin verloren gehen, wenn ...", 1)
add_bullet(doc, "die App deinstalliert oder Website-/Browserdaten gelöscht werden,")
add_bullet(doc, "das Gerät verloren geht, beschädigt wird oder zurückgesetzt werden muss,")
add_bullet(doc, "ein Browserprofil entfernt oder gewechselt wird,")
add_bullet(doc, "das Betriebssystem oder der Browser lokale Daten wegen Speicherproblemen entfernt,")
add_bullet(doc, "eine ältere Sicherung importiert und dadurch neuere lokale Daten ersetzt werden.")

add_heading(doc, "So minimieren Sie das Risiko", 1)
add_step(doc, 1, "Regelmäßig sichern", "Mindestens wöchentlich und zusätzlich nach wichtigen Einträgen eine Sicherung exportieren.", accent=GREEN, fill=GREEN_SOFT)
add_step(doc, 2, "Sicherung getrennt aufbewahren", "Die Sicherungsdatei nicht nur auf demselben Gerät speichern. Nutzen Sie einen geschützten Datenträger oder einen von der Einrichtung freigegebenen Speicherort.", accent=GREEN, fill=GREEN_SOFT)
add_step(doc, 3, "Sicherung schützen", "Exportdateien und PDFs enthalten sensible Gesundheitsdaten und sind nicht zusätzlich verschlüsselt. Gerät und Ablageort mit Passwort beziehungsweise Displaysperre schützen.", accent=GREEN, fill=GREEN_SOFT)
add_step(doc, 4, "Vor Änderungen exportieren", "Vor Deinstallation, Browserbereinigung, Gerätewechsel oder Zurücksetzen des Geräts immer eine aktuelle Sicherung erstellen.", accent=GREEN, fill=GREEN_SOFT)

add_callout(
    doc,
    "Merksatz",
    "Installieren macht KLARA bequemer und offline nutzbar. Nur eine aktuelle, getrennt aufbewahrte Export-Sicherung schützt wirksam vor dauerhaftem Datenverlust.",
    GREEN_SOFT,
    GREEN,
)

add_heading(doc, "Bei Fragen", 1)
add_text(doc, "Sprechen Sie mit Ihrer Fachperson, bevor Sie die App löschen, das Gerät wechseln oder eine Sicherung importieren.")

doc.core_properties.title = "KLARA Installations- und Datensicherungsblatt"
doc.core_properties.subject = "Installation, Nutzung auf mehreren Geräten und Schutz vor Datenverlust"
doc.core_properties.author = "PP.rt - Klinik für Psychiatrie und Psychosomatik Reutlingen"
set_picture_alt_text(doc, "Logo der PP.rt - Klinik für Psychiatrie und Psychosomatik Reutlingen")
doc.save(DOCX)
print(DOCX)
